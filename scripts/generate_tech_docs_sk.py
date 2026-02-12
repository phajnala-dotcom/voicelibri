"""
VoiceLibri — Generátor technickej dokumentácie (SLOVENSKÁ VERZIA)
Generuje komplexný Word dokument (~40-50 strán) s architektúrou,
pipeline diagrammi, API referenciou a modulovým prehľadom.
"""
import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
from lxml import etree

# ── helpers ───────────────────────────────────────────────────

def _keep_with_next(paragraph):
    """Set paragraph to keep with next (avoid orphan before page break)."""
    pPr = paragraph._p.get_or_add_pPr()
    etree.SubElement(pPr, qn('w:keepNext'))

def _keep_lines(paragraph):
    """Set paragraph to keep lines together (no mid-paragraph page break)."""
    pPr = paragraph._p.get_or_add_pPr()
    etree.SubElement(pPr, qn('w:keepLines'))

def _keep_table_together(table):
    """Set all rows in table to keep together (no cross-page break)."""
    for row in table.rows:
        tr = row._tr
        trPr = tr.get_or_add_trPr()
        # Remove any existing cantSplit, then add it
        for existing in trPr.findall(qn('w:cantSplit')):
            trPr.remove(existing)
        etree.SubElement(trPr, qn('w:cantSplit'))
    # Also keep-with-next on all cell paragraphs except last row
    for row_idx, row in enumerate(table.rows):
        for cell in row.cells:
            for para in cell.paragraphs:
                pPr = para._p.get_or_add_pPr()
                etree.SubElement(pPr, qn('w:keepLines'))
                if row_idx < len(table.rows) - 1:
                    etree.SubElement(pPr, qn('w:keepNext'))

def h(doc, text, level=1):
    # Empty paragraph BEFORE heading
    spacer_before = doc.add_paragraph()
    spacer_before.paragraph_format.space_before = Pt(0)
    spacer_before.paragraph_format.space_after = Pt(0)
    spacer_before.paragraph_format.line_spacing = Pt(6)
    _keep_with_next(spacer_before)
    # The heading itself
    heading = doc.add_heading(text, level=level)
    _keep_with_next(heading)
    # Empty paragraph AFTER heading
    spacer_after = doc.add_paragraph()
    spacer_after.paragraph_format.space_before = Pt(0)
    spacer_after.paragraph_format.space_after = Pt(0)
    spacer_after.paragraph_format.line_spacing = Pt(6)
    return heading

def p(doc, text, bold=False, italic=False, sz=11):
    para = doc.add_paragraph()
    r = para.add_run(text)
    r.font.size = Pt(sz)
    if bold: r.bold = True
    if italic: r.italic = True
    return para

def bullets(doc, items):
    for item in items:
        doc.add_paragraph(item, style='List Bullet')

def code(doc, text, label=""):
    if label:
        lp = doc.add_paragraph()
        lr = lp.add_run(label)
        lr.font.size = Pt(9); lr.italic = True
        lr.font.color.rgb = RGBColor(100,100,100)
        _keep_with_next(lp)
    lines = text.strip().split('\n')
    for i, line in enumerate(lines):
        cp = doc.add_paragraph()
        cp.paragraph_format.space_before = Pt(0)
        cp.paragraph_format.space_after  = Pt(0)
        cr = cp.add_run(line)
        cr.font.name = 'Consolas'
        cr.font.size = Pt(8)
        cr.font.color.rgb = RGBColor(40,40,40)
        _keep_lines(cp)
        # Keep all code lines together (keep-with-next except last line)
        if i < len(lines) - 1:
            _keep_with_next(cp)

def tbl(doc, headers, rows):
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.style = 'Light Grid Accent 1'
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i,hdr in enumerate(headers):
        c = t.rows[0].cells[i]; c.text = hdr
        for pp in c.paragraphs:
            for rr in pp.runs: rr.bold = True; rr.font.size = Pt(9)
    for ri,row in enumerate(rows):
        for ci,val in enumerate(row):
            c = t.rows[ri+1].cells[ci]; c.text = str(val)
            for pp in c.paragraphs:
                for rr in pp.runs: rr.font.size = Pt(9)
    # Keep table together across pages
    _keep_table_together(t)
    return t

# ── document ──────────────────────────────────────────────────
doc = Document()

# ── TITULNÁ STRANA ────────────────────────────────────────────
doc.add_paragraph()
title = doc.add_heading('VoiceLibri', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub = doc.add_heading('Technická dokumentácia', level=1)
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
p(doc, 'Verzia 1.2  •  Február 2026', sz=13).alignment = WD_ALIGN_PARAGRAPH.CENTER
p(doc, 'Komerčná AI platforma na dramatizované audioknihy s viacerými hlasmi', italic=True, sz=11).alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph()
p(doc, 'Proprietárna technológia • Pre interné použitie', sz=10, italic=True).alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_page_break()

# ── OBSAH ─────────────────────────────────────────────────────
h(doc, 'Obsah', 1)
toc_items = [
    '1. Účel a rozsah dokumentu',
    '2. Systémová architektúra',
    '   2.1. Vysokoúrovňový diagram',
    '   2.2. Technologický stack',
    '   2.3. Adresárová štruktúra monorepa',
    '3. Audiobook generačný pipeline',
    '   3.1. Prehľad fáz pipeline',
    '   3.2. Fáza 1 — Import a spracovanie e-knihy',
    '   3.3. Fáza 2 — Preklad kapitol',
    '   3.4. Fáza 3 — Extrakcia postáv (CharacterRegistry)',
    '   3.5. Fáza 4 — Hybridná dramatizácia',
    '   3.6. Fáza 5 — Two-Speaker chunking',
    '   3.7. Fáza 6 — TTS syntéza',
    '   3.8. Fáza 7 — Konsolidácia a soundscape',
    '4. Backend moduly — podrobný prehľad',
    '   4.1. index.ts — API server a orchestrácia',
    '   4.2. bookChunker.ts — Parsovanie a chunking',
    '   4.3. hybridDramatizer.ts — Hybridná dramatizácia',
    '   4.4. hybridTagger.ts — Detektor dialógov',
    '   4.5. characterRegistry.ts — Register postáv',
    '   4.6. llmCharacterAnalyzer.ts — LLM analýza postáv',
    '   4.7. voiceAssigner.ts — Priraďovanie hlasov',
    '   4.8. geminiVoices.ts — Databáza hlasov',
    '   4.9. ttsClient.ts — TTS klient',
    '   4.10. twoSpeakerChunker.ts — 2-speaker chunking',
    '   4.11. chapterChunker.ts — Kapitolové chunkovanie',
    '   4.12. chapterTranslator.ts — Preklad kapitol',
    '   4.13. tempChunkManager.ts — Správa temp súborov',
    '   4.14. audiobookManager.ts — Správa knižnice',
    '   4.15. audiobookWorker.ts — Background worker',
    '   4.16. costTracker.ts — Sledovanie nákladov',
    '   4.17. promptConfig.ts — Konfigurácia promptov',
    '   4.18. soundscapeIntegration.ts — Zvukové efekty',
    '   4.19. formatExtractors.ts — Multi-formátové extraktory',
    '   4.20. textCleaner.ts — Čistenie textu',
    '   4.21. audioUtils.ts — Audio utility',
    '   4.22. Ostatné moduly',
    '5. REST API rozhranie',
    '   5.1–5.6. Endpointy podľa kategórií',
    '6. Systém hlasov a postavy',
    '7. Frontend PWA architektúra',
    '   7.1. State management (Zustand)',
    '   7.2. Audio prehrávanie — dva módy',
    '   7.3. Obrazovky a komponenty',
    '8. Mobilná aplikácia (Expo / React Native)',
    '9. Dátové modely a úložisko',
    '10. Výkonnostné obmedzenia a limity',
    '11. Spracovanie chýb a odolnosť',
    '12. Sledovanie nákladov (CostTracker)',
    '13. Testovanie a kvalita',
    '14. Konfigurácia a nasadenie',
    '15. Rozšíriteľnosť a budúce funkcie',
]
for item in toc_items:
    doc.add_paragraph(item)
doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 1. ÚČEL A ROZSAH
# ══════════════════════════════════════════════════════════════
h(doc, '1. Účel a rozsah dokumentu', 1)
p(doc, 'Tento dokument poskytuje komplexnú technickú dokumentáciu platformy VoiceLibri — komerčného AI systému na premenu elektronických kníh na dramatizované audioknihy s viacerými hlasmi. Dokumentácia je určená pre vývojárov na úrovni junior+ a pokrýva celú architektúru, generačný pipeline, API rozhranie, frontend a mobilnú aplikáciu.')
p(doc, 'Cieľová skupina:', bold=True)
bullets(doc, [
    'Junior a senior vývojári pracujúci na VoiceLibri',
    'Noví členovia tímu, ktorí potrebujú onboarding',
    'Architekti hodnotiacie technologické rozhodnutia',
    'QA inžinieri testujúci pipeline a API',
])
p(doc, 'Rozsah dokumentácie:', bold=True)
bullets(doc, [
    'Celková systémová architektúra (monorepo, backend, web PWA, mobilná appka)',
    'Kompletný audiobook generačný pipeline (7 fáz, od importu e-knihy po WAV výstup)',
    'Detailný popis všetkých ~22 backend modulov a ich funkcií',
    'REST API referencia (30+ endpointov)',
    'Frontend architektúra (Zustand stores, TanStack Query, React komponenty, hooky)',
    'Mobilná aplikácia (Expo SDK 54, React Native)',
    'Systém 30 Gemini hlasov a sémantické priraďovanie postavám',
    'Error handling, výkonnostné limity, nákladová analýza',
    'Konfigurácia, nasadenie, testovanie',
])

# ══════════════════════════════════════════════════════════════
# 2. SYSTÉMOVÁ ARCHITEKTÚRA
# ══════════════════════════════════════════════════════════════
h(doc, '2. Systémová architektúra', 1)
h(doc, '2.1. Vysokoúrovňový diagram', 2)
p(doc, 'VoiceLibri je TypeScript monorepo s npm workspaces. Pozostáva z troch aplikácií a zdieľaného súborového úložiska audiokníh.')
code(doc, """
SYSTÉMOVÁ ARCHITEKTÚRA — VYSOKOÚROVŇOVÝ DIAGRAM

┌─────────────────────────────────────────────────────────────────────────┐
│                        MONOREPO  (npm workspaces)                       │
│                                                                         │
│  ┌────────────────────┐  ┌────────────────────┐  ┌───────────────────┐  │
│  │   apps/backend/    │  │    apps/pwa-v2/     │  │   apps/mobile/    │  │
│  │  Express + TS      │  │  React 18 + Vite    │  │  Expo SDK 54      │  │
│  │  Port 3001         │  │  Port 5180          │  │  iOS / Android    │  │
│  │                    │  │                     │  │                   │  │
│  │ ┌────────────────┐ │  │ ┌─────────────────┐ │  │ ┌───────────────┐ │  │
│  │ │ AI Pipeline    │ │  │ │ Zustand Stores  │ │  │ │ Zustand +     │ │  │
│  │ │ ● Gemini LLM   │ │  │ │ ● player        │ │  │ │ AsyncStorage  │ │  │
│  │ │ ● Gemini TTS   │◄├──┤ │ ● library       │ │  │ │ TanStack Q    │ │  │
│  │ │ ● ffmpeg       │ │  │ │ ● theme         │ │  │ │ expo-router   │ │  │
│  │ └────────────────┘ │  │ └─────────────────┘ │  │ └───────────────┘ │  │
│  │                    │  │                     │  │                   │  │
│  │ ┌────────────────┐ │  │ ┌─────────────────┐ │  │ ┌───────────────┐ │  │
│  │ │ 22 modulov     │ │  │ │ TanStack Query  │ │  │ │ Natívne       │ │  │
│  │ │ (src/*.ts)     │ │  │ │ (server cache)  │ │  │ │ komponenty    │ │  │
│  │ └────────────────┘ │  │ └─────────────────┘ │  │ └───────────────┘ │  │
│  └─────────┬──────────┘  └────────────────────┘  └───────────────────┘  │
│            │                                                             │
│            ▼                                                             │
│  ┌───────────────────────────────────────────────────────────────┐       │
│  │                audiobooks/  (súborové úložisko)               │       │
│  │                                                               │       │
│  │  {BookTitle}/                                                 │       │
│  │  ├── metadata.json         (stav, kapitoly, pozícia, prefs)  │       │
│  │  ├── character_registry.json  (postavy, hlasy, aliasy)       │       │
│  │  ├── cost_summary.json     (tokeny, USD náklady)             │       │
│  │  ├── temp/                 (sub-chunk WAV dočasné súbory)    │       │
│  │  │   ├── subchunk_001_000.wav                                │       │
│  │  │   ├── subchunk_001_001.wav                                │       │
│  │  │   └── ...                                                 │       │
│  │  ├── 01_Kapitola_Prvá.wav  (konsolidovaná kapitola)          │       │
│  │  ├── 02_Kapitola_Druhá.wav                                   │       │
│  │  └── ...                                                     │       │
│  └───────────────────────────────────────────────────────────────┘       │
│                                                                         │
│  EXTERNÉ SLUŽBY:                                                        │
│  ┌──────────────────────────────────────────────────────────────────┐   │
│  │  Google Cloud Vertex AI                                          │   │
│  │  ● Gemini 2.5 Flash        → LLM (analýza, preklad, tagging)   │   │
│  │  ● Gemini 2.5 Flash TTS    → Text-to-Speech (max 2 hlasy/req)  │   │
│  └──────────────────────────────────────────────────────────────────┘   │
│                                                                         │
│  ┌──────────────────────────────────────────────────────────────────┐   │
│  │  ffmpeg (lokálny)  → Audio miešanie, soundscape, music intro     │   │
│  └──────────────────────────────────────────────────────────────────┘   │
└─────────────────────────────────────────────────────────────────────────┘
""", label="Diagram 1: Systémová architektúra VoiceLibri")

h(doc, '2.2. Technologický stack', 2)
tbl(doc, ['Komponent', 'Technológia', 'Účel'], [
    ['Backend runtime', 'Node.js + TypeScript + Express', 'REST API server, pipeline orchestrácia'],
    ['LLM engine', 'Google Vertex AI — Gemini 2.5 Flash', 'Analýza postáv, dramatizácia, preklad'],
    ['TTS engine', 'Gemini 2.5 Flash TTS', 'Multi-speaker syntéza (max 2 hlasy/req)'],
    ['Web frontend', 'React 18 + Vite + TypeScript', 'Progresívna webová aplikácia (PWA)'],
    ['Stav (web)', 'Zustand + TanStack Query', 'Lokálny stav + server cache'],
    ['Štýly (web)', 'Tailwind CSS + Neumorphism UI', 'Responzívne 3D-efekt UI'],
    ['Mobilná app', 'React Native + Expo SDK 54', 'Natívna iOS/Android'],
    ['Navigácia (mobile)', 'expo-router v6', 'Súborový routing'],
    ['Stav (mobile)', 'Zustand + AsyncStorage', 'Perzistentný stav'],
    ['Audio', 'WAV (PCM 24kHz, 16-bit, mono)', 'Bezstratový výstup'],
    ['EPUB', 'adm-zip + fast-xml-parser', 'Extrahovanie textu z EPUB'],
    ['PDF', 'pdf-parse', 'Extrahovanie z digitálnych PDF'],
    ['DOCX', 'mammoth', 'Extrahovanie z Word'],
    ['Zvuky', 'ffmpeg (externý)', 'Soundscape mix + music intro'],
    ['Testy', 'vitest', 'Unit testy (backend)'],
])

h(doc, '2.3. Adresárová štruktúra monorepa', 2)
code(doc, """
ebook-reader/                        # Koreň monorepa
├── package.json                     # npm workspaces config
├── apps/
│   ├── backend/                     # ── EXPRESS API SERVER ──
│   │   ├── package.json
│   │   ├── tsconfig.json
│   │   ├── assets/                  # Vstupné e-knihy
│   │   └── src/
│   │       ├── index.ts             # Hlavný server (3260 riadkov)
│   │       ├── bookChunker.ts       # Parsovanie, metadáta, kapitoly (1165 r.)
│   │       ├── hybridDramatizer.ts  # Hybridná dramatizácia
│   │       ├── hybridTagger.ts      # Rule-based tagger (628 r.)
│   │       ├── llmCharacterAnalyzer.ts  # LLM analýza postáv (723 r.)
│   │       ├── characterRegistry.ts # Per-chapter register (654 r.)
│   │       ├── voiceAssigner.ts     # Priraďovanie hlasov
│   │       ├── geminiVoices.ts      # 30 hlasov + trait matching
│   │       ├── ttsClient.ts         # Vertex AI TTS klient (499 r.)
│   │       ├── twoSpeakerChunker.ts # 2-speaker limiter (404 r.)
│   │       ├── chapterChunker.ts    # Ramp-up chunking
│   │       ├── chapterTranslator.ts # Preklad 18 jazykov
│   │       ├── tempChunkManager.ts  # Temp súbory + TTS gen (1834 r.)
│   │       ├── audiobookManager.ts  # Knižnica + metadáta (540 r.)
│   │       ├── audiobookWorker.ts   # Background worker (384 r.)
│   │       ├── costTracker.ts       # Sledovanie nákladov
│   │       ├── promptConfig.ts      # Centrálne prompty + konštanty
│   │       ├── soundscapeIntegration.ts # ffmpeg soundscape (696 r.)
│   │       ├── formatExtractors.ts  # 12+ formátov (744 r.)
│   │       ├── textCleaner.ts       # Regex čistenie (418 r.)
│   │       ├── audioUtils.ts        # WAV concat + silence (84 r.)
│   │       ├── dramatizedProcessor.ts  # PoC orchestrátor
│   │       ├── dramatizedChunkerSimple.ts  # SPEAKER: parser
│   │       ├── dialogueParserSimple.ts # České úvodzovky
│   │       └── parallelPipelineManager.ts # Reset stavu
│   │
│   ├── pwa-v2/                      # ── REACT PWA ──
│   │   └── src/
│   │       ├── App.tsx              # Router + providermi
│   │       ├── screens/             # Library, Generate, Classics, Settings
│   │       ├── components/          # AppShell, FullPlayer, BookItem, ...
│   │       ├── hooks/               # useAudioPlayback, useProgressiveAudio
│   │       ├── stores/              # playerStore, libraryStore, themeStore
│   │       ├── services/api.ts      # HTTP klient (352 r.)
│   │       └── types/               # TypeScript rozhrania
│   │
│   └── mobile/                      # ── EXPO MOBILNÁ APPKA ──
│       ├── app/                     # expo-router routes
│       │   ├── _layout.tsx          # Root + provideri
│       │   ├── (tabs)/              # Explore, Library, Settings
│       │   ├── book/[id].tsx        # Detail knihy
│       │   └── player.tsx           # Modálny prehrávač
│       └── src/
│           ├── components/ui/       # Natívne UI
│           ├── stores/              # Zustand + AsyncStorage
│           ├── services/            # API klient
│           └── theme/               # Dark/light téma
│
├── audiobooks/                      # Generované audioknihy
├── soundscape/                      # Zvukové efekty + hudba
└── docs/                            # Dokumentácia
""", label="Diagram 2: Kompletná adresárová štruktúra")

doc.add_page_break()
# ══════════════════════════════════════════════════════════════
# 3. AUDIOBOOK GENERAČNÝ PIPELINE
# ══════════════════════════════════════════════════════════════
h(doc, '3. Audiobook generačný pipeline', 1)
p(doc, 'Jadrom VoiceLibri je sofistikovaný multi-krokový pipeline, ktorý transformuje e-knihu na dramatizovanú audioknihu. Pipeline beží na pozadí (startBackgroundDramatization) a umožňuje prehrávanie od prvej hotovej kapitoly (progressive playback). Celý process je asynchrónny a non-blocking voči API serveru.')

h(doc, '3.1. Prehľad fáz pipeline', 2)
code(doc, """
AUDIOBOOK GENERAČNÝ PIPELINE — KOMPLETNÝ FLOW

╔══════════════════════════════════════════════════════════════════════════╗
║                     POST /api/book/select                              ║
║  { filename, narratorVoice, targetLanguage, dramatize }                ║
╚══════════════════════════════════════╤═══════════════════════════════════╝
                                       │
                    ┌──────────────────▼──────────────────┐
                    │  FÁZA 1: IMPORT A SPRACOVANIE       │
                    │  loadBookFile()                      │
                    │  ● Detekcia formátu (.epub/.txt/...) │
                    │  ● Extrakcia textu (formatExtractors)│
                    │  ● Čistenie textu (textCleaner)      │
                    │  ● Detekcia kapitol                  │
                    │  ● Parsovanie metadát (autor, jazyk) │
                    │  Výstup: BOOK_CHAPTERS[] (1-based)   │
                    └──────────────────┬──────────────────┘
                                       │
                            ┌──────────▼──────────┐
                            │  200 OK → klient     │
                            │  (book info + chunks) │
                            └──────────┬──────────┘
                                       │
         ┌─────────────────────────────▼──────────────────────────────┐
         │           startBackgroundDramatization()                    │
         │           (beží na pozadí, non-blocking)                    │
         │                                                             │
         │   ╔══ PRE KAŽDÚ KAPITOLU (sekvenčne, 1 → N) ════════════╗  │
         │   ║                                                       ║  │
         │   ║  ┌───────────────────────────────────────────────────┐║  │
         │   ║  │ FÁZA 2: PREKLAD (ak targetLanguage ≠ null)       │║  │
         │   ║  │ ChapterTranslator.translateChapter()              │║  │
         │   ║  │ ● Gemini 2.5 Flash, temperature 0.2              │║  │
         │   ║  │ ● Max 65536 výstupných tokenov                   │║  │
         │   ║  │ ● Normalizácia úvodzoviek po preklade            │║  │
         │   ║  │ ● 18 podporovaných jazykov                       │║  │
         │   ║  └──────────────────────┬────────────────────────────┘║  │
         │   ║                         │                             ║  │
         │   ║  ┌──────────────────────▼────────────────────────────┐║  │
         │   ║  │ FÁZA 3: EXTRAKCIA POSTÁV                         │║  │
         │   ║  │ CharacterRegistry.extractFromChapter()            │║  │
         │   ║  │ ● LLM volanie → JSON: meno, pohlavie, vlastnosti │║  │
         │   ║  │ ● Alias detekcia (sameAs pole)                   │║  │
         │   ║  │ ● Zamknutie hlasu po 1. priradení (locked=true)  │║  │
         │   ║  │ ● BookInfo z kapitol 1-2 (žáner, tón, éra)       │║  │
         │   ║  │ ● Narrator inštrukcia z BookInfo                 │║  │
         │   ║  └──────────────────────┬────────────────────────────┘║  │
         │   ║                         │                             ║  │
         │   ║  ┌──────────────────────▼────────────────────────────┐║  │
         │   ║  │ FÁZA 4: HYBRIDNÁ DRAMATIZÁCIA                    │║  │
         │   ║  │ tagChapterHybrid()                                │║  │
         │   ║  │                                                   │║  │
         │   ║  │  hasDialogue()?──NIE──► NARRATOR: celý_text ($0) │║  │
         │   ║  │       │                                           │║  │
         │   ║  │      ÁNO                                         │║  │
         │   ║  │       │                                           │║  │
         │   ║  │  applyRuleBasedTagging()                          │║  │
         │   ║  │       │                                           │║  │
         │   ║  │  confidence ≥ 85%?                                │║  │
         │   ║  │       │                                           │║  │
         │   ║  │  (Vždy LLM pre speechStyle directives)           │║  │
         │   ║  │       │                                           │║  │
         │   ║  │  LLM fallback → Gemini 2.5 Flash                 │║  │
         │   ║  │  ● extractDialogueParagraphs() (len dialógy)     │║  │
         │   ║  │  ● mergeWithNarration() → finálny text           │║  │
         │   ║  │                                                   │║  │
         │   ║  │  Výstup: "SPEAKER: text" + speechStyle directives│║  │
         │   ║  └──────────────────────┬────────────────────────────┘║  │
         │   ║                         │                             ║  │
         │   ║  ┌──────────────────────▼────────────────────────────┐║  │
         │   ║  │ FÁZA 5: TWO-SPEAKER CHUNKING                     │║  │
         │   ║  │ chunkForTwoSpeakers()                             │║  │
         │   ║  │ ● Max 2 unikátni hovoriaci na chunk               │║  │
         │   ║  │ ● Max 2500 bytov (hard limit API: 4000)           │║  │
         │   ║  │ ● Nepreruší vetu v strede                         │║  │
         │   ║  │ ● Merge consecutive same-speaker segments         │║  │
         │   ║  │ Výstup: TwoSpeakerChunk[] (sub-chunky)           │║  │
         │   ║  └──────────────────────┬────────────────────────────┘║  │
         │   ║                         │                             ║  │
         │   ║  ┌──────────────────────▼────────────────────────────┐║  │
         │   ║  │ FÁZA 6: TTS SYNTÉZA                               │║  │
         │   ║  │ generateSubChunksParallel()                       │║  │
         │   ║  │ ● Paralelizmus: 1 (kap.1), 3 (ostatné)           │║  │
         │   ║  │ ● ttsClient.synthesizeMultiSpeaker()              │║  │
         │   ║  │ ● Retry 3× s exponenciálnym backoffom             │║  │
         │   ║  │ ● Timeout: 120s per požiadavka                    │║  │
         │   ║  │ ● Výstup: WAV buffer (24kHz, 16-bit, mono)       │║  │
         │   ║  │ ● Uloženie: temp/subchunk_CCC_SSS.wav            │║  │
         │   ║  └──────────────────────┬────────────────────────────┘║  │
         │   ║                         │                             ║  │
         │   ║  ┌──────────────────────▼────────────────────────────┐║  │
         │   ║  │ FÁZA 7a: KONSOLIDÁCIA KAPITOLY                   │║  │
         │   ║  │ consolidateChapterFromSubChunks()                 │║  │
         │   ║  │ ● Konkatenácia WAV sub-chunkov                   │║  │
         │   ║  │ ● Tiché pauzy (500ms) medzi sub-chunkmi          │║  │
         │   ║  │ ● Výstup: 01_Chapter_Title.wav                   │║  │
         │   ║  │ ● Aktualizácia metadata.json                     │║  │
         │   ║  └──────────────────────┬────────────────────────────┘║  │
         │   ║                         │                             ║  │
         │   ║  ┌──────────────────────▼────────────────────────────┐║  │
         │   ║  │ FÁZA 7b: SOUNDSCAPE (voliteľná)                  │║  │
         │   ║  │ applySoundscapeToChapter()                        │║  │
         │   ║  │ ● Ambient zvuky mixované cez ffmpeg               │║  │
         │   ║  │ ● Music intro s voice-over narráciou              │║  │
         │   ║  │ ● Audio ducking počas voice-over                  │║  │
         │   ║  │ ● Výstup: {chapter}_soundscape.wav                │║  │
         │   ║  └──────────────────────────────────────────────────┘ ║  │
         │   ║                                                       ║  │
         │   ╚═══════════════════════════════════════════════════════╝  │
         │                                                             │
         │   ● Aktualizácia generationStatus → 'completed'             │
         │   ● Uloženie cost_summary.json                              │
         └─────────────────────────────────────────────────────────────┘
""", label="Diagram 3: Kompletný generačný pipeline (všetkých 7 fáz)")

h(doc, '3.2. Fáza 1 — Import a spracovanie e-knihy', 2)
p(doc, 'Funkcia loadBookFile() v index.ts spracováva vstupný súbor podľa jeho prípony. Výstupom je text rozdelený na kapitoly.', bold=True)
p(doc, 'Podporované formáty (12+):')
tbl(doc, ['Formát', 'Prípony', 'Extraktor', 'Modul'], [
    ['EPUB', '.epub', 'extractTextFromEpub() + extractEpubChapters()', 'bookChunker.ts'],
    ['Text', '.txt', 'fs.readFileSync() + detectTextChapters()', 'bookChunker.ts'],
    ['PDF', '.pdf', 'extractTextFromPdf() (len digitálne PDF, nie sken)', 'formatExtractors.ts'],
    ['HTML', '.html, .htm', 'extractTextFromHtml()', 'formatExtractors.ts'],
    ['MOBI/KF8', '.mobi, .azw, .azw3', 'extractTextFromMobi()', 'formatExtractors.ts'],
    ['Word', '.docx', 'extractTextFromDocx() (mammoth)', 'formatExtractors.ts'],
    ['LibreOffice', '.odt', 'extractTextFromOdt()', 'formatExtractors.ts'],
    ['RTF', '.rtf', 'extractTextFromRtf()', 'formatExtractors.ts'],
    ['Markdown', '.md', 'extractTextFromMarkdown()', 'formatExtractors.ts'],
    ['Pages', '.pages', 'extractTextFromPages()', 'formatExtractors.ts'],
    ['WPS', '.wps', 'extractTextFromWps()', 'formatExtractors.ts'],
    ['DOC (legacy)', '.doc', 'extractTextFromDoc()', 'formatExtractors.ts'],
])
p(doc, 'Detekcia kapitol:')
bullets(doc, [
    'EPUB: Používa OPF manifest spine order, každý spine item = kapitola',
    'TXT: Heuristický parser — hľadá patterny "Chapter X", "KAPITOLA", "Kapitel", rímske číslice',
    'Ak sa kapitoly nenájdu: celý text = 1 kapitola (createSingleChapter)',
    'Text cleaning: removePageNumbers(), removeTOC(), removeEditorialNotes() pred chunk delením',
])

h(doc, '3.3. Fáza 2 — Preklad kapitol', 2)
p(doc, 'ChapterTranslator trieda (chapterTranslator.ts) prekladá kapitoly ak targetLanguage ≠ null.')
tbl(doc, ['Parameter', 'Hodnota'], [
    ['LLM model', 'Gemini 2.5 Flash'],
    ['Temperature', '0.2 (nízka = presný preklad)'],
    ['Max output tokenov', '65 536'],
    ['Retry', '2 opakovania s exponenciálnym backoffom'],
    ['Post-processing', 'normalizeQuotesForDramatization() — curly → straight'],
])
p(doc, 'Podporované jazyky (18):')
p(doc, 'cs (čeština), sk (slovenčina), en (angličtina), de (nemčina), fr (francúzština), es (španielčina), it (taliančina), pt (portugalčina), pl (poľština), ru (ruština), uk (ukrajinčina), nl (holandčina), sv (švédčina), da (dánčina), no (nórčina), fi (fínčina), hu (maďarčina), ro (rumunčina).')

h(doc, '3.4. Fáza 3 — Extrakcia postáv (CharacterRegistry)', 2)
p(doc, 'Per-chapter extrakcia cez LLM (characterRegistry.ts):')
bullets(doc, [
    'LLM prompt obsahuje kompletný zoznam 30 Gemini hlasov → LLM priamo vyberie vhodný hlas',
    'Výstup: JSON s poliami name, gender, traits[], suggestedVoice, aliases[], ageRange, role',
    'Alias detekcia: sameAs pole spája rôzne formy mena (napr. "Pan Harker" → "Jonathan Harker")',
    'Voice locking: po prvom priradení sa hlas nikdy nezmení (locked: true)',
    'BookInfo extrakcia z kap. 1-2: genre (horror, romance...), tone (dark, humorous...), voiceTone, period',
    'BookInfo sa zamkne po 2. kapitole (bookInfoLocked: true)',
    'Narrator TTS inštrukcia: auto-generovaná z BookInfo ("Narrate in a deep, atmospheric tone...")',
    'Uloženie: character_registry.json v audiobook priečinku',
])

h(doc, '3.5. Fáza 4 — Hybridná dramatizácia', 2)
p(doc, 'Modul hybridDramatizer.ts implementuje cost-optimized 3-stratégiový prístup:', bold=True)
code(doc, """
ROZHODOVACÍ STROM HYBRIDNEJ DRAMATIZÁCIE

                    Vstup: text kapitoly
                             │
                    ┌────────▼────────┐
                    │ hasDialogue()?   │  (regex: „..." "..." «...» ...)
                    └────────┬────────┘
                    NIE      │     ÁNO
                    │        │      │
            ┌───────▼──┐   ┌▼──────▼───────────────┐
            │ Stratégia │   │ Stratégia 2 / 3        │
            │ 1: auto   │   │                        │
            │ NARRATOR  │   │ applyRuleBasedTagging()│
            │ $0, 100%  │   │ ● Czech speech verbs   │
            │ hotové    │   │ ● Attribution patterns │
            └──────────┘   │ ● Pronoun analysis     │
                            └────────┬───────────────┘
                                     │
                            ┌────────▼────────┐
                            │ confidence ≥ 85%?│
                            └────────┬────────┘
                            ÁNO      │     NIE
                            │        │      │
               ┌────────────┼────────┼──────┘
               │ (Ale vždy LLM pre  │
               │  speechStyle!)      │
               ▼                     ▼
        ┌──────────────────────────────────────┐
        │ LLM fallback (Gemini 2.5 Flash)      │
        │ ● extractDialogueParagraphs()         │
        │   → len odseky s úvodzovkami          │
        │ ● LLM taguje dialógy + speechStyle    │
        │ ● mergeWithNarration() → finálny text  │
        │ Cena: ~$0.01–0.04/kapitola             │
        └──────────────────────────────────────┘

VÝSTUP FORMÁT (Gemini TTS multi-speaker):
  [Read in a deep, dramatic voice:]
  NARRATOR: Temný hrad sa črtil nad údolím...
  [Whisper fearfully:]
  JONATHAN: Čo to bolo za zvuk?
  NARRATOR: Spýtal sa, trasúc sa na celom tele.
""", label="Diagram 4: Hybridná dramatizácia — rozhodovací strom")

p(doc, 'Kľúčové funkcie hybridTagger.ts:')
tbl(doc, ['Funkcia', 'Účel', 'Detail implementácie'], [
    ['hasDialogue(text)', 'Detekcia dialógov', 'Regex: „…", "…", «…», \'…\', curly quotes'],
    ['countDialogues(text)', 'Počet dialógov', 'Regex counting všetkých typov úvodzoviek'],
    ['applyRuleBasedTagging(text, chars)', 'Rule-based tagger', 'Speech verbs (řekl, zeptala se, 85+), attribution'],
    ['calculateConfidence(tagged, chars)', 'Spoľahlivosť', 'Pomer priradených vs nepriradených dialógov'],
    ['extractDialogueParagraphs(text)', 'Extrakcia dialógov', 'Filter odsekov s úvodzovkami → LLM input'],
    ['mergeWithNarration(orig, tagged, chars)', 'Merge tagging', 'Neoznačený text → NARRATOR:'],
    ['inferGender(name, context)', 'Inferencia pohlavia', '5 metód: CZ koncovky (-ová, -ská), mená, pronomená, slovesá, adjektíva'],
])

h(doc, '3.6. Fáza 5 — Two-Speaker chunking', 2)
p(doc, 'Gemini TTS API podporuje maximálne 2 hovorcov na API volanie. Modul twoSpeakerChunker.ts zabezpečuje dodržanie tohto limitu:')
bullets(doc, [
    'Max 2 unikátni hovoriaci na chunk (napr. NARRATOR + DRACULA, ale nie NARRATOR + DRACULA + MINA)',
    'Max 2500 bytov na chunk (pracovný limit; hard limit API je 4000 B s rezervou pre directives)',
    'Nepreruší vetu v strede — splitSegmentAtSentence() rozdeľuje na vetné hranice',
    'Merge consecutive same-speaker segments — eliminácia krátkych (<50B) segmentov toho istého hovoriaceho',
    'Výstup: TwoSpeakerChunk[] (pole sub-chunkov, každý s max 2 hovorcami a správnou veľkosťou)',
    'formatForMultiSpeakerTTS() — formátuje chunk pre Gemini TTS API volanie',
])

h(doc, '3.7. Fáza 6 — TTS syntéza (ttsClient.ts)', 2)
p(doc, 'TTS klient komunikuje s Google Vertex AI REST API:')
tbl(doc, ['Mód', 'Metóda', 'Použitie', 'Max hovorcov'], [
    ['Single-speaker', 'synthesizeText()', 'Len narrator alebo 1 postava', '1'],
    ['Multi-speaker', 'synthesizeMultiSpeaker()', 'Dramatizovaný text s 2 hovorcami', '2'],
])
p(doc, 'Vlastnosti:')
bullets(doc, [
    'Model: gemini-2.5-flash-tts (konfigurovateľný cez TTS_MODEL env)',
    'Retry: 3 opakovania s exponenciálnym backoffom (2s → 4s → 8s)',
    'Timeout: 120 sekúnd (AbortSignal.timeout)',
    'Safety filter recovery: retry pri intermitentných SAFETY blokoch',
    'Speech style directives: "[Read in a whisper:]", "[Thought, internal monologue:]"',
    'Language code: voliteľný pre krátke texty (prevencia misdetekcie jazyka)',
    'Výstup: WAV buffer (PCM 24kHz, 16-bit, mono)',
    'Voice lookup: 5-úrovňový matching (exact → normalized → case-insensitive → partial → surname)',
])

h(doc, '3.8. Fáza 7 — Konsolidácia a soundscape', 2)
p(doc, 'Po vygenerovaní všetkých sub-chunkov kapitoly:')
p(doc, '7a) Konsolidácia (tempChunkManager.ts → consolidateChapterFromSubChunks):', bold=True)
bullets(doc, [
    'Načítanie všetkých subchunk_CCC_SSS.wav z temp/ priečinka',
    'Zoradenie podľa číselného indexu',
    'WAV konkatenácia s tichými pauzami (500ms, SUBCHUNK_GAP_MS)',
    'audioUtils.ts: concatenateWavBuffers() + addSilence()',
    'Výstup: 01_Chapter_Title.wav',
    'Aktualizácia metadata.json: isConsolidated = true, duration v sekundách',
])
p(doc, '7b) Soundscape (soundscapeIntegration.ts, voliteľná):', bold=True)
bullets(doc, [
    'Ambient vrstva: environmentálne zvuky (príroda, mesto, interiér)',
    'Mixovanie cez ffmpeg: amix filter s amerge + loudnorm normalizáciou',
    'Music intro: krátke intro (5-15s) s voice-over ("Túto audioknihu pre vás pripravil VoiceLibri")',
    'Voice-over TTS v cieľovom jazyku knihy',
    'Music ducking: hudba sa stlmí na -12dB počas voice-over',
    'Theme selection: getSoundscapeThemeOptions() — skórovanie tém podľa kľúčových slov v texte',
    'Výstup: {chapter}_soundscape.wav',
])

doc.add_page_break()
# ══════════════════════════════════════════════════════════════
# 4. BACKEND MODULY
# ══════════════════════════════════════════════════════════════
h(doc, '4. Backend moduly — podrobný prehľad', 1)
p(doc, 'Backend pozostáva z ~22 TypeScript modulov v apps/backend/src/. Táto sekcia popisuje účel, exportované funkcie a vzájomné väzby každého modulu.')

# 4.1
h(doc, '4.1. index.ts — API server a orchestrácia (3260 riadkov)', 2)
p(doc, 'Hlavný Express server. Obsahuje všetky API endpointy a orchestruje celý background pipeline.')
p(doc, 'Hlavné zodpovednosti:', bold=True)
bullets(doc, [
    'Express server s CORS, JSON body parsing (50MB limit), statickým servírovaním audiobooks/',
    'Všetky REST API endpointy (30+ endpointov — viď sekcia 5)',
    'loadBookFile() — detekcia formátu, extrakcia textu, parsovanie metadát, detekcia kapitol',
    'startBackgroundDramatization() — orchestrácia pipeline per kapitola (preklad → postavy → dramatizácia → TTS → konsolidácia → soundscape)',
    'In-memory globálny stav: BOOK_CHAPTERS[], CHAPTER_SUBCHUNKS Map, VOICE_MAP, NARRATOR_VOICE, TARGET_LANGUAGE',
    'Audio serving s prioritou: subchunk file → consolidated chapter → legacy temp → memory cache → 202 not ready',
    'Automatická konsolidácia: checkAndConsolidateReadyChapters() po každom sub-chunk played',
    'AbortController pre cancellation background procesu pri výbere novej knihy',
])
p(doc, 'Kľúčový in-memory stav:', bold=True)
code(doc, """
// Globálne premenné v index.ts
let BOOK_TEXT: string = '';
let BOOK_CHAPTERS: Chapter[] = [];          // 1-based pole kapitol
let BOOK_METADATA: BookMetadata | null;
let BOOK_FORMAT: string = '';
let CURRENT_BOOK_FILE: string = '';
let VOICE_MAP: Record<string, string> = {}; // postava → Gemini hlas
let NARRATOR_VOICE: string = 'Enceladus';   // default narrator
let TARGET_LANGUAGE: string | null = null;
let COST_TRACKER: CostTracker | null = null;

const CHAPTER_SUBCHUNKS = new Map<number, TwoSpeakerChunk[]>();
const CHAPTER_DRAMATIZED = new Map<number, string>();
const CHAPTER_DRAMATIZATION_LOCK = new Map<number, Promise<string>>();
let TOTAL_SUBCHUNKS = 0;
const audioCache = new Map<string, Buffer>();
""")

# 4.2
h(doc, '4.2. bookChunker.ts — Parsovanie a chunking (1165 riadkov)', 2)
p(doc, 'Zodpovednosť: Načítanie e-kníh, parsovanie metadát, extrakcia kapitol, chunking textu.')
tbl(doc, ['Exportovaná funkcia', 'Účel'], [
    ['chunkBookText(text, options)', 'Rozdelenie textu na chunky (byte limit, sentence boundaries)'],
    ['parseBookMetadata(filename, content, format)', 'Strategy pattern → BookMetadata (autor, jazyk, názov)'],
    ['extractEpubChapters(filepath)', 'EPUB → pole kapitol (OPF spine order)'],
    ['detectTextChapters(text)', 'TXT → heuristická detekcia kapitol (Chapter, KAPITOLA...)'],
    ['createSingleChapter(text, title)', 'Fallback: celý text = 1 kapitola'],
    ['getBookInfo()', 'Vracia aktuálne BookMetadata'],
    ['formatDuration(seconds)', 'Formátovanie trvania (HH:MM:SS)'],
])

# 4.3
h(doc, '4.3. hybridDramatizer.ts — Hybridná dramatizácia', 2)
p(doc, 'Zodpovednosť: Cost-optimized dramatizácia (60–80% úspora vs. čistý LLM).')
tbl(doc, ['Exportovaná funkcia', 'Účel'], [
    ['tagChapterHybrid(text, chars, bookInfo)', 'Hlavná funkcia — 3-stratégiový rozhodovací strom'],
    ['dramatizeBookStreaming(chapters, ...)', 'AsyncGenerator — yielding dramatizované kapitoly postupne'],
    ['dramatizeFirstChapterHybrid(text, chars)', 'Špeciálna cesta pre 1. kapitolu (rýchly time-to-audio)'],
])
p(doc, 'Tri stratégie:')
bullets(doc, [
    'Stratégia 1 (bez dialógov): Auto-NARRATOR tag, cena $0, confidence 100%',
    'Stratégia 2 (vysoká confidence ≥85%): Rule-based tagging, cena $0 (ale LLM pre speechStyle)',
    'Stratégia 3 (nízka confidence): LLM fallback len na dialógové odseky, cena ~$0.01–0.04/kap.',
])

# 4.4
h(doc, '4.4. hybridTagger.ts — Detektor dialógov (628 riadkov)', 2)
p(doc, 'Zodpovednosť: Rule-based detekcia a taggovanie dialógov, inferencia pohlavia.')
tbl(doc, ['Exportovaná funkcia', 'Účel'], [
    ['hasDialogue(text)', 'Prítomnosť dialógu (regex: 7+ typov úvodzoviek)'],
    ['countDialogues(text)', 'Počet dialógov v texte'],
    ['applyRuleBasedTagging(text, characters)', 'Rule-based speaker attribution (85+ Czech speech verbs)'],
    ['calculateConfidence(taggedText, characters)', 'Skóre spoľahlivosti (0.0–1.0)'],
    ['extractDialogueParagraphs(text)', 'Filtrovanie odsekov s dialógmi pre LLM'],
    ['mergeWithNarration(originalText, taggedDialogues, chars)', 'Zlúčenie LLM tagov s narráciou'],
    ['inferGender(name, contextText)', 'Inferencia pohlavia: CZ koncovky, mená, pronomená, slovesá, adjektíva'],
])

# 4.5
h(doc, '4.5. characterRegistry.ts — Register postáv (654 riadkov)', 2)
p(doc, 'Trieda CharacterRegistry — kumulatívny stav postáv naprieč kapitolami.')
tbl(doc, ['Metóda / Export', 'Účel'], [
    ['extractFromChapter(chapterText, chapterNum)', 'LLM extrakcia postáv z kapitoly'],
    ['getVoiceMap()', 'Record<string, string> — postava → Gemini hlas'],
    ['getAllCharacters()', 'Všetky postavy s locked/unlocked stavom'],
    ['saveToFile(dir)', 'Uloženie do character_registry.json'],
    ['loadFromFile(dir)', 'Načítanie registra z disku'],
    ['BookInfo', 'Interface: genre, tone, voiceTone, period'],
])

# 4.6
h(doc, '4.6. llmCharacterAnalyzer.ts — LLM analýza (723 riadkov)', 2)
p(doc, 'Trieda GeminiCharacterAnalyzer — Gemini LLM volania pre analýzu postáv.')
tbl(doc, ['Metóda / Export', 'Účel'], [
    ['analyzeFullBook(text)', 'Celá kniha → CharacterProfile[] (meno, pohlavie, vlastnosti, hlas, vek)'],
    ['tagChapterWithVoices(chapterText, chars)', 'Tagging kapitoly s voice markers + speechStyle'],
    ['toTTSSpeakerAlias(name)', 'Normalizácia: "Jan Novák" → "JANNOVAK" (bez diakritiky, uppercase)'],
    ['CharacterProfile interface', 'name, gender, traits[], suggestedVoice, aliases[], ageRange, role'],
])

# 4.7
h(doc, '4.7. voiceAssigner.ts — Priraďovanie hlasov', 2)
p(doc, 'Heuristické priraďovanie hlasov na základe pohlavia, veku a trait clusterov. Zabezpečuje unikátnosť.')
tbl(doc, ['Funkcia', 'Účel'], [
    ['assignVoices(characters, narratorVoice)', 'Priradenie unikátnych Gemini hlasov'],
    ['saveVoiceMap(dir, voiceMap)', 'Uloženie voice_map.json'],
    ['loadVoiceMap(dir)', 'Načítanie voice mapy'],
])

# 4.8
h(doc, '4.8. geminiVoices.ts — Databáza hlasov', 2)
p(doc, '30 predefinovaných Gemini TTS hlasov so sémantickým matchingom.')
tbl(doc, ['Kategória', 'Počet', 'Príklady (Gemini → alias)'], [
    ['Mužské (low)', '5', 'Algieba→Albert (deep, authoritative), Alnilam→Milan (warm)'],
    ['Mužské (medium)', '7', 'Puck→Peter (youthful), Achird→Arthur (storyteller)'],
    ['Mužské (high)', '4', 'Umbriel→Urban (energetic), Laomedeia→Leo (clear)'],
    ['Ženské (low)', '4', 'Gacrux→Grace (strong), Vindemiatrix→Viola (theatrical)'],
    ['Ženské (medium)', '6', 'Achernar→Ash (professional), Sulafat→Sarah (soothing)'],
    ['Ženské (high)', '4', 'Zephyr→Zara (light), Erinome→Erin (sweet)'],
])
p(doc, 'selectVoiceForCharacter(traits, gender, ageRange): Sémantické skórovanie s 27 trait clustermi. Podporuje české aj anglické traits (napr. "babička" → elderly cluster).')

# 4.9
h(doc, '4.9. ttsClient.ts — TTS klient (499 riadkov)', 2)
p(doc, 'Google Vertex AI Gemini TTS klient.')
tbl(doc, ['Metóda', 'Účel', 'Detail'], [
    ['synthesizeText(text, voiceName, opts)', 'Single-speaker TTS', 'Voliteľný speechStyle, languageCode'],
    ['synthesizeMultiSpeaker(turns, voiceConfig)', 'Multi-speaker TTS', 'Max 2 hlasy, multiSpeakerVoiceConfig'],
])
p(doc, 'Konfigurácia: model=gemini-2.5-flash-tts, maxOutputTokens=8192, endpoint=us-central1-aiplatform.googleapis.com')

# 4.10
h(doc, '4.10. twoSpeakerChunker.ts — 2-speaker chunking (404 riadkov)', 2)
tbl(doc, ['Funkcia', 'Účel'], [
    ['chunkForTwoSpeakers(dramatizedText)', 'Rozdelenie na TwoSpeakerChunk[] (max 2 hovoriaci, max 2500B)'],
    ['formatForMultiSpeakerTTS(chunk)', 'Formátovanie pre Gemini TTS API'],
    ['getUniqueSpeakers(chunk)', 'Zoznam unikátnych hovorcov v chunku'],
    ['splitSegmentAtSentence(segment, maxBytes)', 'Rozdelenie na vetné hranice'],
])

# 4.11
h(doc, '4.11. chapterChunker.ts — Kapitolové chunkovanie', 2)
p(doc, 'Ramp-up stratégia: progresívne zvyšovanie veľkosti chunkov pre rýchly time-to-first-audio.')
code(doc, """
Ramp-up sekvencia (byte limity per chunk):
  Chunk 1:  300 B   ← veľmi malý = rýchle audio
  Chunk 2:  500 B
  Chunk 3:  800 B
  Chunk 4: 1200 B
  Chunk 5: 1800 B
  Chunk 6: 2500 B
  Chunk 7+: 3500 B  ← plný limit (pod 4000B hard limit)

GEMINI_TTS_HARD_LIMIT = 4000 bytov (validácia)
""")

# 4.12
h(doc, '4.12. chapterTranslator.ts — Preklad kapitol', 2)
tbl(doc, ['Export', 'Účel'], [
    ['ChapterTranslator class', 'Preklad kapitol cez Gemini 2.5 Flash'],
    ['translateChapter(text, targetLang)', 'Preloženie jednej kapitoly'],
    ['normalizeQuotesForDramatization(text)', 'Normalizácia úvodzoviek po preklade'],
    ['SUPPORTED_LANGUAGES', '18 jazykov s BCP-47 kódmi'],
])

# 4.13
h(doc, '4.13. tempChunkManager.ts — Správa temp súborov (1834 riadkov)', 2)
p(doc, 'Najväčší modul. Zodpovedný za TTS generáciu, temp caching, voice lookup a konsolidáciu.')
tbl(doc, ['Funkcia', 'Účel'], [
    ['generateAndSaveTempChunk(chunk, index, ...)', 'Generácia jedného sub-chunk WAV a uloženie do temp/'],
    ['generateSubChunksParallel(chapter, chunks, ...)', 'Paralelná TTS generácia (1-3 concurrent)'],
    ['consolidateChapterFromSubChunks(bookTitle, ch)', 'Konsolidácia sub-chunkov → kapitola WAV'],
    ['consolidateChapterSmart(bookTitle, ch)', 'Smart konsolidácia s kontrolou completeness'],
    ['lookupVoice(speakerName, voiceMap)', '5-úrovňový voice matching'],
    ['startPreDramatization(chapterText, ...)', 'Look-ahead dramatizácia v pozadí'],
    ['dramatizeTextCore(text, voiceMap, ...)', 'Core dramatizácia + chunking'],
    ['toBCP47(langCode)', 'Konverzia jazykového kódu na BCP-47'],
])
p(doc, 'Voice lookup priority (lookupVoice):')
bullets(doc, [
    '1. Exact match: voiceMap["Jonathan Harker"]',
    '2. Normalized: trimmed, lowercased key match',
    '3. Case-insensitive: porovnanie bez ohľadu na veľkosť písmen',
    '4. Partial match: kľúč obsahuje meno hovoriaceho alebo naopak',
    '5. Surname match: posledné slovo mena sa zhoduje',
    'Fallback: narrator voice (ak sa nič nenájde)',
])

# 4.14
h(doc, '4.14. audiobookManager.ts — Správa knižnice (540 riadkov)', 2)
tbl(doc, ['Funkcia', 'Účel'], [
    ['createAudiobookFolder(title)', 'Vytvorenie {audiobooks}/{title}/ + temp/'],
    ['saveAudiobookMetadata(title, metadata)', 'Uloženie metadata.json'],
    ['loadAudiobookMetadata(title)', 'Načítanie metadát'],
    ['listAudiobooks()', 'Zoznam všetkých audiokníh'],
    ['deleteAudiobook(title)', 'Zmazanie priečinka + metadát'],
    ['getSubChunkPath(title, chapter, subchunk)', 'Cesta k temp/subchunk_CCC_SSS.wav'],
    ['countChapterSubChunks(title, chapter)', 'Počet sub-chunkov pre kapitolu'],
    ['isChapterConsolidated(title, chapter)', 'Kontrola či existuje konsolidovaný WAV'],
])

# 4.15
h(doc, '4.15. audiobookWorker.ts — Background worker (384 riadkov)', 2)
p(doc, 'Trieda AudiobookGenerationWorker (extends EventEmitter) pre queue-based generáciu.')
tbl(doc, ['Metóda', 'Účel'], [
    ['addBook(title, config)', 'Pridanie knihy do fronty'],
    ['processQueue()', 'Spracovanie fronty (FIFO)'],
    ['generateAudiobook(job)', 'Generácia jednej audioknihy'],
    ['generateAllChunks(job)', 'Paralelná generácia chunkov (max 2 concurrent)'],
    ['consolidateAllChapters(job)', 'Konsolidácia všetkých kapitol'],
    ['getProgress()', 'Aktuálny stav generácie'],
])

# 4.16
h(doc, '4.16. costTracker.ts — Sledovanie nákladov', 2)
p(doc, 'Sleduje tokeny a USD per audiokniha cez 4 fázy pipeline.')
tbl(doc, ['Fáza', 'Input $/M tokenov', 'Output $/M tokenov', 'Typická cena/kap.'], [
    ['Extrakcia postáv', '$0.15', '$0.60', '$0.002–0.005'],
    ['Preklad', '$0.15', '$0.60', '$0.01–0.05'],
    ['Dramatizácia', '$0.30', '$2.50', '$0.01–0.04'],
    ['Audio generácia', '$0.15', '$0.60', '$0.01–0.03'],
])
p(doc, 'Token estimation: slová × koeficient (slovanské jazyky: 2.15, angličtina: 1.38, default: 1.76). Výstup: cost_summary.json.')

# 4.17
h(doc, '4.17. promptConfig.ts — Konfigurácia promptov', 2)
p(doc, 'Single source of truth pre všetky LLM prompty, temperatures a konštanty.')
tbl(doc, ['Export', 'Účel', 'Hodnota / Detail'], [
    ['getCharacterExtractionPrompt()', 'Prompt pre extrakciu postáv', 'Obsahuje všetkých 30 hlasov'],
    ['getVoiceTaggingPrompt()', 'Prompt pre voice tagging', 'SPEAKER: formát + speechStyle'],
    ['getTranslationPrompt()', 'Prompt pre preklad', 'Zachovanie formátovania, úvodzoviek'],
    ['buildNarratorInstruction(bookInfo)', 'Narrator TTS style', '"Narrate in a [tone]..."'],
    ['DEFAULT_NARRATOR_VOICE', 'Default narrator', 'Enceladus'],
    ['SUBCHUNK_GAP_MS', 'Pauza medzi sub-chunkmi', '500 ms'],
    ['DRAMATIZATION_TIMEOUT_MS', 'Timeout pre dramatizáciu', '300 000 ms (5 min)'],
])

# 4.18
h(doc, '4.18. soundscapeIntegration.ts (696 riadkov)', 2)
tbl(doc, ['Funkcia', 'Účel'], [
    ['applySoundscapeToChapter(chapterPath, theme)', 'Ambient mix cez ffmpeg'],
    ['createMusicIntro(bookTitle, voiceText, ...)', 'Hudobné intro s voice-over narráciou'],
    ['getSoundscapeThemeOptions(text)', 'Skórovanie tém podľa kľúčových slov'],
    ['duckAudio(music, voiceover, outputPath)', 'Music ducking (-12dB) počas voice-over'],
])

# 4.19
h(doc, '4.19. formatExtractors.ts (744 riadkov)', 2)
p(doc, 'Multi-formátové extraktory textu. Každý extraktor vracia čistý text.')
tbl(doc, ['Funkcia', 'Knižnica', 'Poznámka'], [
    ['extractTextFromEpub(path)', 'adm-zip + fast-xml-parser', 'OPF spine order'],
    ['extractTextFromPdf(path)', 'pdf-parse', 'Len digitálne PDF (nie OCR sken)'],
    ['extractTextFromDocx(path)', 'mammoth', 'Zachováva štruktúru odsekov'],
    ['extractTextFromHtml(path)', 'regex/cheerio', 'Strip HTML tagov'],
    ['extractTextFromMobi(path)', 'custom parser', 'MOBI/KF8/AZW'],
    ['extractTextFromOdt(path)', 'adm-zip (ZIP → content.xml)', 'OpenDocument formát'],
    ['extractTextFromRtf(path)', 'regex stripping', 'RTF control words removal'],
    ['extractTextFromMarkdown(path)', 'regex', 'Strip markdown syntaxu'],
    ['SUPPORTED_EXTENSIONS', 'Set', '.epub, .txt, .pdf, .html, .mobi, .docx, .odt, .rtf, .md, .pages, .wps, .doc'],
    ['SUPPORTED_MIME_TYPES', 'Map', 'MIME → handler mapping'],
])

# 4.20-4.22
h(doc, '4.20. textCleaner.ts (418 riadkov)', 2)
p(doc, 'Regex-based čistenie textu od artefaktov e-kníh.')
tbl(doc, ['Funkcia', 'Účel'], [
    ['cleanText(text)', 'Orchestrátor — volá všetky čistenia'],
    ['removePageNumbers(text)', 'Odstránenie čísiel strán'],
    ['removeTOC(text)', 'Odstránenie obsahu (Table of Contents)'],
    ['removeEditorialNotes(text)', 'Odstránenie redakčných poznámok'],
    ['removePublisherInfo(text)', 'Odstránenie informácií o vydavateľstve'],
    ['normalizeWhitespace(text)', 'Normalizácia medzier a riadkov'],
])

h(doc, '4.21. audioUtils.ts (84 riadkov)', 2)
tbl(doc, ['Funkcia', 'Účel'], [
    ['concatenateWavBuffers(buffers)', 'Konkatenácia viacerých WAV bufferov do jedného'],
    ['addSilence(durationMs)', 'Generovanie tichej pauzy (24kHz, 16-bit, mono)'],
])

h(doc, '4.22. Ostatné moduly', 2)
tbl(doc, ['Modul', 'Riadkov', 'Účel'], [
    ['dramatizedProcessor.ts', '~200', 'PoC pipeline orchestrátor (staršia verzia)'],
    ['dramatizedChunkerSimple.ts', '~150', 'Parser "SPEAKER: text" formátu'],
    ['dialogueParserSimple.ts', '~100', 'Jednoduchý detektor dialógov (české úvodzovky „...")'],
    ['parallelPipelineManager.ts', '~80', 'Reset globálneho stavu pri prepínaní kníh'],
])

doc.add_page_break()
# ══════════════════════════════════════════════════════════════
# 5. REST API
# ══════════════════════════════════════════════════════════════
h(doc, '5. REST API rozhranie', 1)
p(doc, 'Backend exponuje REST API na porte 3001. Všetky endpointy sú definované v index.ts.')

h(doc, '5.1. Health & knihy', 2)
tbl(doc, ['Metóda', 'Endpoint', 'Popis', 'Odpoveď'], [
    ['GET', '/api/health', 'Health check', '200: {status, bookLoaded}'],
    ['GET', '/api/books', 'Zoznam kníh v assets/', '200: [{filename, size, type}]'],
    ['POST', '/api/book/select', 'Výber + inicializácia', '200: {title, author, chapters, chunks}'],
    ['GET', '/api/book/info', 'Metadáta aktuálnej knihy', '200: BookMetadata'],
    ['GET', '/api/book/consolidated', 'Stav konsolidácie', '200: {chapters: [{index, consolidated}]}'],
    ['POST', '/api/book/from-text', 'Nová kniha z textu/base64', '200: {title, chapters}'],
    ['POST', '/api/book/from-url', 'Nová kniha z URL', '200: {title, chapters}'],
])

h(doc, '5.2. TTS & audio', 2)
tbl(doc, ['Metóda', 'Endpoint', 'Popis', 'Odpoveď'], [
    ['POST', '/api/tts/chunk', 'Audio pre sub-chunk', '200: WAV / 202: {retryAfterMs}'],
    ['GET', '/api/audiobooks/:title/subchunks/:ch/:sub', 'Stream sub-chunk', '200: WAV stream'],
    ['GET', '/api/audiobooks/:title/chapters/:ch', 'Konsolidovaná kapitola', '200: WAV stream'],
])
p(doc, 'Hlavičky v odpovedi POST /api/tts/chunk:', bold=True)
bullets(doc, [
    'X-Cache: subchunk_file | chapter_file | legacy_temp | file_scan | memory_cache | generated',
    'X-Is-Whole-Chapter: true/false — indikuje konsolidovanú kapitolu',
    'X-Seek-Offset-Sec: float — offset pre navigáciu v konsolidovanej kapitole',
    'Content-Type: audio/wav',
])

h(doc, '5.3. Knižnica', 2)
tbl(doc, ['Metóda', 'Endpoint', 'Popis'], [
    ['GET', '/api/audiobooks', 'Zoznam audiokníh'],
    ['GET', '/api/audiobooks/:title', 'Metadáta konkrétnej audioknihy'],
    ['DELETE', '/api/audiobooks/:title', 'Zmazanie audioknihy'],
    ['POST', '/api/audiobooks/generate', 'Background generácia (worker)'],
    ['GET', '/api/audiobooks/:title/progress', 'Priebeh generácie'],
    ['GET', '/api/audiobooks/worker/status', 'Stav workera'],
])

h(doc, '5.4. Pozícia & preferencie', 2)
tbl(doc, ['Metóda', 'Endpoint', 'Popis'], [
    ['PUT', '/api/audiobooks/:title/position', 'Uloženie pozície prehrávania'],
    ['GET', '/api/audiobooks/:title/position', 'Načítanie pozície'],
    ['PUT', '/api/audiobooks/:title/preferences', 'Uloženie preferencií (hlas, rýchlosť, soundscape)'],
    ['GET', '/api/audiobooks/:title/preferences', 'Načítanie preferencií'],
])

h(doc, '5.5. Dramatizácia & soundscape', 2)
tbl(doc, ['Metóda', 'Endpoint', 'Popis'], [
    ['GET', '/api/dramatization/status', 'Stav background dramatizácie (phase, chapter, timeout)'],
    ['GET', '/api/dramatize/check/:bookFile', 'Kontrola cache pre dramatizáciu'],
    ['POST', '/api/dramatize/process', 'Spracovanie pred-tagovaného textu'],
    ['GET', '/api/dramatize/voice-map', 'Voice map pre dramatizovaný text'],
    ['GET', '/api/audiobooks/:title/soundscape/themes', 'Dostupné soundscape témy'],
])

h(doc, '5.6. Error formát', 2)
code(doc, """
// Štandardný error response (všetky endpointy):
{
  "error": "ERROR_CODE",         // napr. "NO_BOOK_LOADED", "CHUNK_NOT_READY"
  "message": "Ľudsky čitateľný popis chyby"
}

// HTTP kódy:
// 200 — OK (audio alebo JSON)
// 202 — Accepted, not ready yet (retryAfterMs pre polling)
// 400 — Bad request (chýbajúce parametre, nepodporovaný formát)
// 404 — Not found (audiokniha, kapitola, sub-chunk)
// 500 — Internal server error (TTS zlyhanie, file system error)
""")

doc.add_page_break()
# ══════════════════════════════════════════════════════════════
# 6. SYSTÉM HLASOV
# ══════════════════════════════════════════════════════════════
h(doc, '6. Systém hlasov a postavy', 1)

h(doc, '6.1. Databáza 30 Gemini hlasov', 2)
p(doc, 'Každý hlas má: name (Gemini API meno), alias (frontend meno), gender, pitch (low/medium/high), characteristic (popis štýlu).')
tbl(doc, ['Gemini meno', 'Alias', 'Pohlavie', 'Výška', 'Charakteristika'], [
    ['Algieba', 'Albert', 'M', 'low', 'Deep, authoritative narrator'],
    ['Alnilam', 'Milan', 'M', 'low', 'Warm, rich baritone'],
    ['Schedar', 'Stefan', 'M', 'low', 'Gravelly, mature'],
    ['Rasalgethi', 'Richard', 'M', 'low', 'Commanding, noble'],
    ['Puck', 'Peter', 'M', 'medium', 'Youthful, enthusiastic'],
    ['Achird', 'Arthur', 'M', 'medium', 'Classic storyteller'],
    ['Orus', 'Oliver', 'M', 'medium', 'Friendly, conversational'],
    ['Zubenelgenubi', 'Zbyněk', 'M', 'medium', 'Thoughtful, measured'],
    ['Sadachbia', 'Samuel', 'M', 'medium', 'Gentle, calming'],
    ['Enceladus', 'Eric', 'M', 'medium', 'Versatile narrator (DEFAULT)'],
    ['Isonoe', 'Ivan', 'M', 'medium', 'Expressive, dynamic'],
    ['Umbriel', 'Urban', 'M', 'high', 'Energetic, bright'],
    ['Laomedeia', 'Leo', 'M', 'high', 'Clear, articulate'],
    ['Despina', 'Daniel', 'M', 'high', 'Light, pleasant'],
    ['Daphne', 'David', 'M', 'high', 'Smooth, refined'],
    ['Fenrir', 'Filip', 'M', 'medium', 'Bold, adventurous'],
    ['Achernar', 'Ash', 'F', 'medium', 'Professional, neutral'],
    ['Gacrux', 'Grace', 'F', 'low', 'Strong, commanding'],
    ['Vindemiatrix', 'Viola', 'F', 'low', 'Theatrical, dramatic'],
    ['Charon', 'Charlotte', 'F', 'low', 'Mysterious, dark'],
    ['Pulcherrima', 'Paula', 'F', 'low', 'Elegant, mature'],
    ['Sulafat', 'Sarah', 'F', 'medium', 'Soothing, warm'],
    ['Leda', 'Linda', 'F', 'medium', 'Friendly, approachable'],
    ['Callirrhoe', 'Clara', 'F', 'medium', 'Intelligent, precise'],
    ['Autonoe', 'Anna', 'F', 'medium', 'Expressive, emotional'],
    ['Elara', 'Elena', 'F', 'medium', 'Soft, gentle'],
    ['Zephyr', 'Zara', 'F', 'high', 'Light, airy, playful'],
    ['Erinome', 'Erin', 'F', 'high', 'Sweet, youthful'],
    ['Aoede', 'Andrea', 'F', 'high', 'Musical, lyrical'],
    ['Kore', 'Karen', 'F', 'high', 'Bright, cheerful'],
])

h(doc, '6.2. Sémantický matching (trait clusters)', 2)
p(doc, 'Funkcia selectVoiceForCharacter() používa 27 sémantických clusterov pre matching:')
code(doc, """
PRÍKLADY TRAIT CLUSTEROV (z geminiVoices.ts):

"elderly"   → ["elderly", "old", "aged", "wise", "ancient", "babička", "dědeček", "grandmother"]
"young"     → ["young", "youthful", "teen", "child", "boy", "girl", "dítě", "mladý"]
"villain"   → ["villain", "evil", "sinister", "dark", "wicked", "cruel", "zlý", "temný"]
"romantic"  → ["romantic", "passionate", "loving", "tender", "gentle", "láskyplný"]
"military"  → ["military", "soldier", "commanding", "stern", "disciplined", "vojenský"]
"noble"     → ["noble", "aristocratic", "royal", "dignified", "regal", "šlechtický"]
"mysterious"→ ["mysterious", "enigmatic", "secretive", "cryptic", "záhadný"]
"comedic"   → ["comedic", "funny", "humorous", "witty", "sarcastic", "vtipný"]
...

Scoring: Každý trait → najlepší cluster → zoznam vhodných hlasov
         Hlas s najvyšším skóre vyhráva (penalizácia za duplicity)
""")

doc.add_page_break()
# ══════════════════════════════════════════════════════════════
# 7. FRONTEND PWA
# ══════════════════════════════════════════════════════════════
h(doc, '7. Frontend PWA architektúra', 1)
p(doc, 'Web frontend je React 18 PWA (Progressive Web App) s Vite build systémom, Tailwind CSS štýlmi a neumorphism dizajnom.')

h(doc, '7.1. State management (Zustand)', 2)
code(doc, """
FRONTEND STATE — 3 ZUSTAND STORES

┌─────────────────────────────────────────────────────────────────┐
│ playerStore  (voicelibri-player)  — 313 riadkov                │
│ Perzistencia: localStorage                                      │
├─────────────────────────────────────────────────────────────────┤
│ Stav:                                                           │
│   currentBook: Audiobook | null                                 │
│   currentChapterIndex: number                                   │
│   currentTime: number (sekundy v kapitole)                      │
│   isPlaying: boolean                                            │
│   playbackSpeed: number (0.5 – 2.0)                            │
│   volume: number (0 – 1)                                        │
│   progressivePlayback: {                                        │
│     enabled: boolean,                                           │
│     isGenerating: boolean,                                      │
│     currentSubChunk: number,                                    │
│     totalSubChunks: number,                                     │
│     phase: 'loading' | 'playing' | 'done'                      │
│   }                                                             │
│   sleepTimer: { endTime, duration, active }                     │
│ Akcie:                                                          │
│   play(), pause(), seek(time), nextChapter(), prevChapter()     │
│   setPlaybackSpeed(speed), setVolume(vol)                       │
│   startProgressivePlayback(book), stopProgressivePlayback()     │
│   setSleepTimer(minutes)                                        │
└─────────────────────────────────────────────────────────────────┘

┌─────────────────────────────────────────────────────────────────┐
│ libraryStore  (voicelibri-library)                              │
├─────────────────────────────────────────────────────────────────┤
│ Stav: books[], searchQuery, sortBy, filterStatus                │
│       generationProgress: Record<string, number>                │
│ Akcie: addBook(), removeBook(), updateProgress()                │
│ Computed: filteredBooks() — search + filter + sort pipeline      │
└─────────────────────────────────────────────────────────────────┘

┌─────────────────────────────────────────────────────────────────┐
│ themeStore  (voicelibri-theme)                                  │
├─────────────────────────────────────────────────────────────────┤
│ Stav: theme ('light' | 'dark' | 'system')                      │
│ Akcie: setTheme(), toggleTheme()                                │
│ Efekt: automatická detekcia system preference                   │
└─────────────────────────────────────────────────────────────────┘
""", label="Diagram 6: Frontend Zustand stores")

h(doc, '7.2. Audio prehrávanie — dva módy', 2)
tbl(doc, ['Mód', 'Hook', 'Veľkosť', 'Princíp'], [
    ['Chapter mode', 'useAudioPlayback', '180 riadkov', 'Štandardný HTML5 Audio, prehrávanie konsolidovaných WAV kapitol, auto-advance na ďalšiu, ukladanie pozície každých 10s cez PUT /api/audiobooks/:title/position'],
    ['Progressive mode', 'useProgressiveAudioPlayback', '350 riadkov', 'Streamovanie sub-chunkov počas generácie, in-memory Blob URL cache, polling ready stavu (3s interval), auto-prechod na chapter mode po konsolidácii kapitoly'],
])
p(doc, 'Frontend automaticky prepína medzi módmi:')
bullets(doc, [
    'Ak existuje konsolidovaná kapitola WAV → chapter mode',
    'Ak sa kapitola ešte generuje → progressive mode (sub-chunk polling)',
    'Po konsolidácii → seamless prechod na chapter mode',
])

h(doc, '7.3. Obrazovky a komponenty', 2)
p(doc, 'Router (React Router v6 v App.tsx):', bold=True)
tbl(doc, ['Route', 'Obrazovka', 'Veľkosť', 'Popis'], [
    ['/', 'LibraryScreen', '247 r.', 'Zoznam audiokníh, search, sort, filter, delete, play'],
    ['/generate', 'GenerateScreen', '504 r.', 'Vytvorenie novej audioknihy: upload súboru, paste textu, URL import, nastavenia hlasu'],
    ['/classics', 'ClassicsScreen', '153 r.', 'Prehľad klasík (placeholder s mock dátami)'],
    ['/settings', 'SettingsScreen', '215 r.', 'Téma, playback preferencie, about'],
])
p(doc, 'Hlavné komponenty:', bold=True)
tbl(doc, ['Komponent', 'Veľkosť', 'Účel'], [
    ['AppShell', '~100 r.', 'Layout wrapper: sidebar/header + content area + bottom nav'],
    ['FullPlayer', '380 r.', 'Plný prehrávač: progress bar, controls, chapter list, speed, sleep timer'],
    ['MiniPlayer', '183 r.', 'Minimalizovaný prehrávač: play/pause, progress, title'],
    ['BookItem', '350 r.', 'Karta audioknihy: cover, title, author, progress, play/delete akcie'],
    ['BookList', '80 r.', 'Grid/list zobrazenie BookItem komponentov'],
    ['BottomNavigation', '~80 r.', 'Spodná navigácia: Library, Generate, Classics, Settings'],
])
p(doc, 'API klient (services/api.ts, 352 riadkov):', bold=True)
p(doc, 'Všetky fetch() volania na localhost:3001/api. Funkcie: fetchBooks(), selectBook(), fetchAudiobooks(), generateFromText(), generateFromUrl(), getSubChunkAudioUrl(), getChapterAudioUrl(), savePosition(), loadPosition(), getDramatizationStatus() atď.')

doc.add_page_break()
# ══════════════════════════════════════════════════════════════
# 8. MOBILNÁ APLIKÁCIA
# ══════════════════════════════════════════════════════════════
h(doc, '8. Mobilná aplikácia (Expo / React Native)', 1)
p(doc, 'Mobilná appka je mobile-only (žiadna web podpora). Komunikuje s rovnakým backendom na porte 3001.')
tbl(doc, ['Aspekt', 'Detail'], [
    ['Framework', 'Expo SDK 54.0.31, React Native 0.81.5, React 19.1.0'],
    ['Navigácia', 'expo-router v6.0.21 (súborový routing: app/(tabs)/, app/book/[id].tsx)'],
    ['State management', 'Zustand 5.0.0 + AsyncStorage (nie localStorage!)'],
    ['Server state', 'TanStack Query (fetchBooks, fetchAudiobook, ...)'],
    ['Animácie', 'moti + react-native-reanimated v4'],
    ['Gestá', 'react-native-gesture-handler + @gorhom/bottom-sheet'],
    ['Téma', 'Custom ThemeContext (dark/light), expo-blur pre glassmorphism'],
    ['Icons', '@expo/vector-icons (Ionicons, MaterialIcons)'],
])
p(doc, 'Provider stack (app/_layout.tsx):', bold=True)
code(doc, """
<GestureHandlerRootView>           // Gesto handling
  <QueryClientProvider>            // TanStack Query
    <ThemeProvider>                 // Dark/light téma
      <StatusBar style="light" />
      <Stack screenOptions={{ headerShown: false }}>
        <Stack.Screen name="(tabs)" />
        <Stack.Screen name="book/[id]" />
        <Stack.Screen name="player" options={{ presentation: 'modal' }} />
      </Stack>
    </ThemeProvider>
  </QueryClientProvider>
</GestureHandlerRootView>
""")
p(doc, 'Tab navigácia:', bold=True)
tbl(doc, ['Tab', 'Ikona', 'Obrazovka'], [
    ['Explore', 'compass', 'Objavovanie audiokníh (odporúčania, žánre)'],
    ['Library', 'library', 'Knižnica stiahnutých/generovaných audiokníh'],
    ['Settings', 'settings', 'Nastavenia (téma, účet, hlas)'],
])
p(doc, 'Mobile Zustand stores:', bold=True)
bullets(doc, [
    'settingsStore: téma, jazyk, defaultný hlas, playback speed (AsyncStorage persist)',
    'bookStore: stiahnuté knihy, offline stav',
    'playerStore: aktuálna kniha, kapitola, pozícia, hlasitosť',
])

doc.add_page_break()
# ══════════════════════════════════════════════════════════════
# 9. DÁTOVÉ MODELY
# ══════════════════════════════════════════════════════════════
h(doc, '9. Dátové modely a úložisko', 1)
h(doc, '9.1. AudiobookMetadata (metadata.json)', 2)
code(doc, """
interface AudiobookMetadata {
  title: string;                    // Názov knihy
  author: string;                   // Autor
  language: string;                 // Kód jazyka (cs, sk, en, ...)
  totalChapters: number;            // Počet kapitol
  chapters: ChapterMetadata[];      // Pole metadát kapitol
  generationStatus: 'not-started' | 'in-progress' | 'completed';
  lastUpdated: string;              // ISO timestamp

  // Voice & dramatization
  voiceMap?: Record<string, string>;  // Postava → Gemini hlas
  sourceFile?: string;
  isDramatized?: boolean;
  dramatizationType?: 'llm-only' | 'hybrid-optimized';
  charactersFound?: number;
  dramatizationCost?: number;

  // Pozícia prehrávania (cross-device sync)
  playback?: {
    currentChapter: number;         // 0-based
    currentTime: number;            // sekundy
    lastPlayedAt: string;           // ISO timestamp
  };

  // Užívateľské preferencie
  userPreferences?: {
    narratorVoice?: string;
    playbackSpeed?: number;         // 0.75 – 2.0
    soundscapeMusicEnabled?: boolean;
    soundscapeAmbientEnabled?: boolean;
    soundscapeThemeId?: string;
  };
}

interface ChapterMetadata {
  index: number;
  title: string;
  filename: string;                 // "01_Chapter_Title.wav"
  duration: number;                 // sekundy
  isGenerated: boolean;
  isConsolidated?: boolean;
  subChunksTotal?: number;
  subChunksGenerated?: number;
}
""", label="Dátový model: AudiobookMetadata")

h(doc, '9.2. Backend in-memory stav', 2)
p(doc, 'Backend udržiava globálny stav v pamäti (nie databáza):')
tbl(doc, ['Premenná', 'Typ', 'Účel'], [
    ['BOOK_TEXT', 'string', 'Celý extrahovaný text aktuálnej knihy'],
    ['BOOK_CHAPTERS', 'Chapter[]', '1-based pole kapitol'],
    ['BOOK_METADATA', 'BookMetadata | null', 'Metadáta (autor, jazyk, ...)'],
    ['VOICE_MAP', 'Record<string, string>', 'Postava → Gemini hlas'],
    ['NARRATOR_VOICE', 'string', 'Aktuálny narrator (default: Enceladus)'],
    ['TARGET_LANGUAGE', 'string | null', 'Cieľový jazyk prekladu'],
    ['CHAPTER_SUBCHUNKS', 'Map<number, TwoSpeakerChunk[]>', 'Kapitola → sub-chunky'],
    ['CHAPTER_DRAMATIZED', 'Map<number, string>', 'Kapitola → dramatizovaný text'],
    ['audioCache', 'Map<string, Buffer>', 'In-memory audio cache'],
    ['COST_TRACKER', 'CostTracker | null', 'Sledovanie nákladov'],
])

doc.add_page_break()
# ══════════════════════════════════════════════════════════════
# 10. LIMITY
# ══════════════════════════════════════════════════════════════
h(doc, '10. Výkonnostné obmedzenia a limity', 1)
tbl(doc, ['Parameter', 'Limit', 'Zdôvodnenie'], [
    ['Max veľkosť e-knihy', '50 MB', 'express.json({ limit: "50mb" })'],
    ['Max bytov na TTS chunk', '4000 B (hard limit)', 'Gemini TTS API limit'],
    ['Pracovný limit na chunk', '2500 B', 'Rezerva pre speechStyle directive'],
    ['Max hovorcov na TTS req', '2', 'Gemini multiSpeakerVoiceConfig limit'],
    ['TTS timeout', '120 s', 'AbortSignal.timeout v fetch()'],
    ['TTS retry', '3×', 'Exponenciálny backoff: 2s, 4s, 8s'],
    ['Dramatization timeout', '5 min / kapitola', 'DRAMATIZATION_TIMEOUT_MS'],
    ['Paralelná TTS gen.', '1 (kap.1), 3 (ostatné)', 'API rate limit balancing'],
    ['Ramp-up chunks', '300→500→800→1200→1800→2500→3500 B', 'Rýchly time-to-first-audio'],
    ['Silence gap', '500 ms', 'SUBCHUNK_GAP_MS'],
    ['Audio formát', 'WAV 24kHz, 16-bit, mono', 'Vertex AI default'],
    ['Subchunk polling', '2-3 s', 'retryAfterMs v 202 response'],
])

doc.add_page_break()
# ══════════════════════════════════════════════════════════════
# 11. ERROR HANDLING
# ══════════════════════════════════════════════════════════════
h(doc, '11. Spracovanie chýb a odolnosť', 1)
tbl(doc, ['Chyba', 'Riešenie', 'Modul'], [
    ['Gemini TTS 500/503', 'Retry 3× s exponenciálnym backoffom', 'ttsClient.ts'],
    ['Gemini Safety block', 'Retry (intermitentný filter)', 'ttsClient.ts'],
    ['TTS timeout (120s)', 'Abort + retry', 'ttsClient.ts'],
    ['Chunk > 4000 B', 'splitSegmentAtSentence → word boundary', 'twoSpeakerChunker.ts'],
    ['Dramatizácia zlyhá', 'Fallback: NARRATOR: celý text', 'hybridDramatizer.ts'],
    ['Preklad zlyhá', 'Pokračovanie s originálnym textom', 'chapterTranslator.ts'],
    ['Voice not found', '5-úrovňový lookupVoice → default narrator', 'tempChunkManager.ts'],
    ['Kapitola bez dialógov', 'Auto-tag NARRATOR ($0)', 'hybridDramatizer.ts'],
    ['Chapter timeout (5 min)', 'phase="failed", continue next', 'index.ts'],
    ['PDF sken (nie OCR)', 'Odmietnutie s error správou', 'formatExtractors.ts'],
    ['Nepodporovaný formát', '400 + zoznam podporovaných', 'index.ts'],
    ['Network error (frontend)', '202 retry loop', 'api.ts'],
    ['Memory exhaustion', 'Clear chunk maps, restart', 'index.ts'],
    ['Background abort', 'AbortController.abort()', 'index.ts'],
])

# ══════════════════════════════════════════════════════════════
# 12. COST TRACKING
# ══════════════════════════════════════════════════════════════
h(doc, '12. Sledovanie nákladov', 1)
p(doc, 'CostTracker sleduje tokeny a USD na 4 fázy pipeline. Token estimation: slová × jazykový koeficient.')
tbl(doc, ['Fáza', 'Input $/M tok.', 'Output $/M tok.', 'Typicky/kap.'], [
    ['Extrakcia postáv', '$0.15', '$0.60', '$0.002–0.005'],
    ['Preklad', '$0.15', '$0.60', '$0.01–0.05'],
    ['Dramatizácia', '$0.30', '$2.50', '$0.01–0.04'],
    ['Audio generácia', '$0.15', '$0.60', '$0.01–0.03'],
])
p(doc, 'Jazykové koeficienty tokenizácie:')
tbl(doc, ['Jazyk', 'Koeficient', 'Príklad'], [
    ['Slovanské (cs, sk, pl, ru, uk)', '2.15', '1000 slov ≈ 2150 tokenov'],
    ['Angličtina', '1.38', '1000 slov ≈ 1380 tokenov'],
    ['Ostatné', '1.76', 'Priemerný odhad'],
])
p(doc, 'Výstup: cost_summary.json v audiobook priečinku s breakdown per fáza a celkovým súčtom.')

# ══════════════════════════════════════════════════════════════
# 13. TESTOVANIE
# ══════════════════════════════════════════════════════════════
h(doc, '13. Testovanie a kvalita', 1)
p(doc, 'Kombinovaný prístup: unit testy (vitest) + manuálne overovanie kvality audia.')
p(doc, 'Automatizované testy:', bold=True)
bullets(doc, [
    'bookChunker.test.ts — byte limity, sentence boundaries, chapter detection, ramp-up sekvencia',
    'hybridTagger.test.ts — rule-based tagging accuracy, confidence calculation',
    'Spustenie: npx vitest (z apps/backend/)',
])
p(doc, 'Manuálne testovanie:', bold=True)
bullets(doc, [
    'Kompletný pipeline: import → dramatizácia → TTS → prehrávanie',
    'Audio quality check: posluch WAV (hlasové priradenie, plynulosť, prirodzenosť)',
    'Multi-formátové testy: EPUB, TXT, PDF, DOCX',
    'Jazykové testy: čeština, angličtina, slovenčina, nemčina',
    'API monitoring: curl na progress, status, health endpointy',
    'Frontend: navigácia, progressive playback, sleep timer, speed control',
])

# ══════════════════════════════════════════════════════════════
# 14. KONFIGURÁCIA
# ══════════════════════════════════════════════════════════════
h(doc, '14. Konfigurácia a nasadenie', 1)
h(doc, '14.1. Environmentálne premenné', 2)
tbl(doc, ['Premenná', 'Popis', 'Default'], [
    ['GOOGLE_CLOUD_PROJECT', 'ID Google Cloud projektu', '(povinné)'],
    ['GOOGLE_CLOUD_LOCATION', 'Región Vertex AI', 'us-central1'],
    ['GOOGLE_APPLICATION_CREDENTIALS', 'Cesta k service account JSON', '(povinné)'],
    ['PORT', 'Port backendu', '3001'],
    ['TTS_MODEL', 'Gemini TTS model', 'gemini-2.5-flash-tts'],
    ['LLM_MODEL', 'Gemini LLM model', 'gemini-2.5-flash'],
])
h(doc, '14.2. Development príkazy', 2)
code(doc, """
# Koreň monorepa:
npm run dev             # Backend (3001) + PWA (5180) súčasne
npm run dev:backend     # Len backend
npm run dev:pwa         # Len PWA
npm run build           # Production build

# Mobilná appka:
cd apps/mobile
npx expo start          # Expo dev server
npx expo start --tunnel # Pre cross-network prístup

# Testy:
cd apps/backend
npx vitest              # Unit testy
npx vitest --watch      # Watch mód
""")

# ══════════════════════════════════════════════════════════════
# 15. ROZŠÍRITEĽNOSŤ
# ══════════════════════════════════════════════════════════════
h(doc, '15. Rozšíriteľnosť a budúce funkcie', 1)
tbl(doc, ['Oblasť', 'Aktuálny stav', 'Plánované'], [
    ['Formáty', '12+ (EPUB, TXT, PDF, HTML, MOBI, DOCX, ODT, RTF, MD, Pages, WPS, DOC)', 'OCR PDF podpora'],
    ['Jazyky', '18 jazykov (preklad + TTS)', 'Ďalšie jazyky, jazyková detekcia'],
    ['Úložisko', 'Súborový systém + metadata.json', 'Cloud storage, PostgreSQL'],
    ['Autentifikácia', 'Single user (žiadna auth)', 'JWT + multi-user'],
    ['Platby', 'Žiadne', 'Stripe/payment integrácia'],
    ['Mobile', 'Základná funkčnosť', 'Offline prehrávanie, download kapitol'],
    ['Soundscape', 'Ambient + music intro (ffmpeg)', 'AI-generované soundscapes'],
    ['Kvalita', 'Manual testing', 'E2E testy (Playwright), CI/CD'],
])

# ── ULOŽIŤ ────────────────────────────────────────────────────
out_dir = r'c:\Users\hajna\VoiceLibri\docs'
os.makedirs(out_dir, exist_ok=True)
out_path = os.path.join(out_dir, 'VoiceLibri_Technical_Documentation_SK.docx')
doc.save(out_path)

# stats
heading_count = len([pp for pp in doc.paragraphs if pp.style.name.startswith('Heading')])
para_count = len(doc.paragraphs)
table_count = len(doc.tables)
print(f"✓ Slovenská dokumentácia uložená: {out_path}")
print(f"  Nadpisy:  {heading_count}")
print(f"  Odseky:   {para_count}")
print(f"  Tabuľky:  {table_count}")
